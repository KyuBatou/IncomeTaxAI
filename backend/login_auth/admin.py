from django.contrib.auth.forms import AdminPasswordChangeForm, UserChangeForm, UserCreationForm
from django.contrib.auth.admin import UserAdmin as BaseUserAdmin
from django.utils.translation import gettext_lazy as _
from django.contrib.auth.forms import UserChangeForm
from django.db.models.functions import Coalesce
from rangefilter.filters import DateRangeFilter
from django.utils.timezone import localtime
from import_export import resources, fields
from django.utils.html import format_html
from datetime import datetime, timedelta
from user_visit.models import UserVisit
from django.utils.timezone import now
from docx.enum.section import WD_ORIENT
from django.http import HttpResponse
from django.db.models import Value
from docx.shared import Pt, Inches
from django.contrib import admin
from django.urls import reverse
from django.db.models import Q
from docx.shared import Pt
from docx import Document
from django import forms
from .models import *
import openpyxl

class CustomDatePickerWidget(forms.DateInput):

    def __init__(self, attrs=None):
        default_attrs = {
            'type': 'date',
            'max': '9999-12-31',
            'autocomplete': 'off',
            'style': (
                'display: inline-block; '
                'width: 18em;'
                'padding: 6px 10px;'
                'margin: 10px 0px;'
                'border: 1px solid #d1d5db;'
                'border-radius: 6px;'
                'background-color: #ffffff;'
                'box-shadow: inset 0 1px 2px rgba(0,0,0,0.1);'
            )
        }
        if attrs:
            default_attrs.update(attrs)
        super().__init__(attrs=default_attrs, format='%Y-%m-%d')

    def format_value(self, value):
        if isinstance(value, str):
            return value
        if value: return value.strftime('%Y-%m-%d')
        return ''

class CustomUserCreationForm(UserCreationForm):
    class Meta:
        model = User
        fields = ('email', 'name', 'company', 'passwordd', 'is_active', 'is_staff', 'is_salesman')

# Custom User Change Form
class CustomUserChangeForm(UserChangeForm):
    class Meta:
        model = User
        fields = ('email', 'name', 'company', 'passwordd', 'is_active', 'is_staff', 'is_salesman')

class SortingOrderFilter(admin.SimpleListFilter):
    title = 'Sorting Order'
    parameter_name = 'sort'

    def lookups(self, request, model_admin):
        return [
            ('state', 'State (A-Z)'),
            ('-state', 'State (Z-A)'),
            ('city', 'City (A-Z)'),
            ('-city', 'City (Z-A)'),
            ('email', 'Email (A-Z)'),
            ('-email', 'Email (Z-A)'),
            ('-user_monthly_page_view', 'Most Page Views'),
            ('user_monthly_page_view', 'Least Page Views'),
            ('-pk', 'Newest First (Default)'),
            ('pk', 'Oldest First'),
        ]

    def queryset(self, request, queryset):
        if self.value():
            if self.value() in ['user_monthly_page_view', '-user_monthly_page_view']:
                return queryset.annotate(pageview_fixed=Coalesce('user_monthly_page_view', Value(0)))
            return queryset
        return queryset

class CityFilter(admin.SimpleListFilter):
    title = 'City'
    parameter_name = 'city'

    def lookups(self, request, model_admin):
        cities = User.objects.values_list('city', flat=True).distinct()
        return [(city, city) for city in cities if city]

    def queryset(self, request, queryset):
        if self.value():
            return queryset.filter(city__iexact=self.value())
        return queryset

class SalesmanFilter(admin.SimpleListFilter):
    title = _('Salesman')
    parameter_name = 'salesman'

    def lookups(self, request, model_admin):
        return [(salesman.id, salesman.name) for salesman in User.objects.filter(is_salesman=True, is_active=True).order_by('name')]

    def queryset(self, request, queryset):
        return queryset.filter(salesman_id=self.value()) if self.value() else queryset

class ValidDateFilter(admin.SimpleListFilter):
    title = _('User Expired')
    parameter_name = 'valid_date'

    def lookups(self, request, model_admin):
        return [
            ('Expired', _('Expired')),
            ('30', _('Next 30 Days')),
            ('60', _('Next 60 Days')),
            ('90', _('Next 90 Days')),
            ('Approved', _('Approved')),
            ('Un Approved', _('Un Approved')),
        ]

    def queryset(self, request, queryset):
        q = Q()
        today = now().date()
        if self.value() == 'Expired': q.add(Q(valid_date__lt=today), Q.AND)
        elif self.value() == 'Approved': q.add(Q(status='Approved') & Q(valid_date__gt=today), Q.AND)
        elif self.value() == 'Un Approved': q.add(Q(status='Pending Approval') | Q(valid_date__isnull=True), Q.AND)
        elif self.value() in ['30', '60', '90']:
            end_date = today + timedelta(days=int(self.value()))
            q.add(Q(valid_date__range=(today, end_date)), Q.AND)
        return queryset.filter(q)

class UserShortResource(resources.ModelResource):
    email = fields.Field(column_name='Email', readonly=True)
    email_name_company = fields.Field(column_name='Email, Name & Company', readonly=True)
    address_all = fields.Field(column_name='Address All', readonly=True)
    founder_member_detail = fields.Field(column_name='Founder Member', readonly=True)

    def dehydrate_email(self, obj):
        return obj.email if obj.email else ''

    def dehydrate_email_name_company(self, obj):
        parts = []
        if obj.email: parts.append(obj.email)
        if obj.name: parts.append(obj.name)
        if obj.company: parts.append(obj.company)
        return ' \n'.join(parts) if parts else ''

    def dehydrate_name_company(self, obj):
        parts = []
        if obj.name: parts.append(obj.name)
        if obj.company: parts.append(obj.company)
        return ' \n'.join(parts)

    def dehydrate_address_all(self, obj):
        parts = []
        if obj.address: parts.append(obj.address)
        if obj.city: parts.append(obj.city)
        if obj.state: parts.append(obj.state)
        if obj.pin: parts.append(str(obj.pin))
        if obj.mobileno: parts.append(str(obj.mobileno))
        if obj.telephone: parts.append(str(obj.telephone))
        return ', '.join(parts)


    def dehydrate_founder_member_detail(self, obj):
        parts = []
        if obj.founder_member == 'Founder Member': parts.append('Founder Member')
        if obj.founder_member_amount: parts.append(f'LAST AMOUNT: {obj.founder_member_amount}')
        if obj.salesman: parts.append(obj.salesman.name)
        return ' \n'.join(parts)

    def dehydrate_regdate(self, obj):
        return obj.regdate.strftime('%d-%m-%Y') if obj.regdate else '' 

    def dehydrate_valid_date(self, obj):
        return obj.valid_date.strftime('%d-%m-%Y') if obj.valid_date else ''

    class Meta:
        model = User
        fields = ('regdate', 'valid_date')
        export_order = ('email_name_company', 'address_all', 'regdate', 'valid_date', 'founder_member_detail')

class UserResource(UserShortResource):

    email_password = fields.Field(column_name='Email & Password', readonly=True)

    def dehydrate_email_password(self, obj):
        return f"{obj.email or ''} \n{obj.passwordd or ''}"

    class Meta:
        model = User
        fields = ('regdate', 'valid_date')
        export_order = ('email_password', 'name_company', 'address_all', 'regdate', 'valid_date', 'founder_member_detail')


@admin.register(User)
class UserAdmin(BaseUserAdmin):
    formfield_overrides = {models.DateField: {'widget': CustomDatePickerWidget},}
    actions = ["export_as_short_word", "export_as_word", "export_as_short_excel", "export_as_excel"]
    form = CustomUserChangeForm
    resource_class = UserResource
    add_form = CustomUserCreationForm
    change_password_form = AdminPasswordChangeForm
    ordering = ('-pk',)

    list_display = ('valid_and_regdate', 'email_and_password', 'name_and_company', 'status', 'mobile_detailed', 'address_detailed', 'founder_detail', 'view_detail', 'status_display', 'delete_link')
    list_display_links = ('valid_and_regdate', 'email_and_password', 'name_and_company',)
    list_filter = (SortingOrderFilter, 'status', CityFilter, 'state', 'founder_member', SalesmanFilter, ValidDateFilter, ('valid_date', DateRangeFilter), )
    search_fields = ('email', 'name', 'city', 'company', 'mobileno')
    readonly_fields = ('passwordd', 'regdate', 'last_login', 'user_monthly_page_view', 'user_page_view_month', 'is_login', 'visit_and_approval_details', 'ai_monthly_page_view', 'ai_page_view_month',)
    sortable_by = ('-pk', 'state', 'city', 'email', 'user_monthly_page_view', 'user_page_view_month')

    fieldsets = (
        (_('Approve Registration'), {
            'fields': (
                ('status',),
                ('is_multi_user',),
                ('valid_date',),
                ('founder_member', 'founder_member_amount',),
                ('salesman',),
                ('monthly_page_view',),
                ('monthly_ai_page_view',),
            ),
        }),
        (_('Log Entry Details'), {
            'fields': (
                ('visit_and_approval_details',),
            ),
        }),

        (_('Personal Info'), {
            'fields': (
                ('email', 'username', 'name', 'company', 'fax', 'address', 'city', 'state', 'pin', 'mobileno', 'telephone', 'passwordd',),
                ('is_active', 'is_staff', 'is_salesman', 'is_future', 'is_ai_web', 'is_superuser',)
            ),
        }),
        (_('Important Details'), {
            'fields': (
                ('regdate', 'last_login', 'user_monthly_page_view', 'user_page_view_month', 'ai_monthly_page_view', 'ai_page_view_month', 'is_login',),
            ),
        }),
    )

    def delete_link(self, obj):
        url = reverse('admin:%s_%s_delete' % (obj._meta.app_label, obj._meta.model_name), args=[obj.pk])
        return format_html('<a class="button" style="color:red;" href="{}">Delete</a>', url)

    delete_link.short_description = 'Delete'
    delete_link.allow_tags = True

    def get_ordering(self, request):
        sort_param = request.GET.get('sort', '-pk')
        return [sort_param]

    def get_queryset(self, request):
        qs = super().get_queryset(request).filter(is_salesman=False, is_staff=False, is_superuser=False)
        return qs

    def get_actions(self, request):
        actions = super().get_actions(request)
        if "delete_selected" in actions:
            del actions["delete_selected"]
        return actions
    
    def export_as_word(modeladmin, request, queryset):
        doc = Document()
        doc.add_heading("User Data Export", level=1)
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)

        resource = UserResource()
        field_names = resource.get_export_order()

        table = doc.add_table(rows=1, cols=len(field_names))
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for idx, field in enumerate(field_names):
            hdr_cells[idx].text = field.replace("_", " ").title()

        for obj in queryset:
            row_cells = table.add_row().cells
            for idx, field in enumerate(field_names):
                if hasattr(resource, f"dehydrate_{field}"):
                    value = getattr(resource, f"dehydrate_{field}")(obj)
                else:
                    value = getattr(obj, field, "")

                cell_para = row_cells[idx].paragraphs[0]
                for i, line in enumerate(str(value).split("\n")):
                    if i > 0:
                        cell_para.add_run().add_break()
                    run = cell_para.add_run(line)
                    run.font.size = Pt(10)

        filename = f"user_data.docx"
        response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response

    export_as_word.short_description = "Download as Word"

    def export_as_short_word(modeladmin, request, queryset):
        doc = Document()
        doc.add_heading("User Data Export", level=1)
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)

        resource = UserShortResource()
        field_names = resource.get_export_order()

        table = doc.add_table(rows=1, cols=len(field_names))
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for idx, field in enumerate(field_names):
            hdr_cells[idx].text = field.replace("_", " ").title()

        for obj in queryset:
            row_cells = table.add_row().cells
            for idx, field in enumerate(field_names):
                if hasattr(resource, f"dehydrate_{field}"):
                    value = getattr(resource, f"dehydrate_{field}")(obj)
                else:
                    value = getattr(obj, field, "")

                cell_para = row_cells[idx].paragraphs[0]
                for i, line in enumerate(str(value).split("\n")):
                    if i > 0:
                        cell_para.add_run().add_break()
                    run = cell_para.add_run(line)
                    run.font.size = Pt(10)

        filename = f"user_short_data.docx"
        response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response

    export_as_short_word.short_description = "Download Short as Word"

    def export_as_excel(modeladmin, request, queryset):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Users"

        columns = ["Name", "Username", "Founder Member", "Founder Member Amount", 'Tsr Name', "id", "Passwordd", "Company", 
            "Address", "City", "State", "Pin", "telephone", "Mobile No", "Fax", "Email", "Status", "User Monthly Page View"]
        ws.append(columns)

        for user in queryset:
            sales_name = user.salesman.name if user.salesman else ''
            founder_member = user.founder_member if user.founder_member == 'Founder Member' else ''
            ws.append([ user.name, user.username, founder_member, user.founder_member_amount, sales_name, user.id, user.passwordd, user.company,
                user.address, user.city, user.state, user.pin, user.telephone, user.mobileno, user.fax, user.email, user.status, user.user_monthly_page_view
            ])

        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        response["Content-Disposition"] = f'attachment; filename="Users.xlsx"'
        wb.save(response)
        
        return response

    export_as_excel.short_description = "Download as Excel"

    def export_as_short_excel(modeladmin, request, queryset):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Users"

        columns = ["Name", "Username", "Founder Member", "Founder Member Amount", 'Tsr Name', "id",  "Company", 
            "Address", "City", "State", "Pin", "telephone", "Mobile No", "Fax", "Email", "Status", "User Monthly Page View"]
        ws.append(columns)

        for user in queryset:
            sales_name = user.salesman.name if user.salesman else ''
            founder_member = user.founder_member if user.founder_member == 'Founder Member' else ''
            ws.append([ user.name, user.username, founder_member, user.founder_member_amount, sales_name, user.id, user.company,
                user.address, user.city, user.state, user.pin, user.telephone, user.mobileno, user.fax, user.email, user.status, user.user_monthly_page_view
            ])

        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        response["Content-Disposition"] = f'attachment; filename="Users.xlsx"'
        wb.save(response)
        
        return response

    export_as_short_excel.short_description = "Download Short as Excel"

    def visit_details(self, obj):
        visits = AuthSession.objects.filter(user=obj)
        if not visits.exists():
            return "No visits"
        rows = ""
        for i, visit in enumerate(visits, start=1):
            rows += (
                f"<tr>"
                f"<td>{i}</td>"
                f"<td>{localtime(visit.created_at).strftime('%d-%m-%Y %I:%M:%S %p')}</td>"
                f"<td>{visit.ip_address}</td>"
                f"<td>{visit.user_agent}</td>"
                f"</tr>"
            )
        
        visit_table = (
            f"<div style='flex: 1; margin-right: 10px;'>"
            f"<h3>Visit Details</h3>"
            f"<div style='max-height: 200px; overflow-y: auto;'>"
            f"<table style='border: 1px solid #ddd; border-collapse: collapse; width: 100%;'>"
            f"<thead>"
            f"<tr style='background-color: #f9f9f9;'>"
            f"<th>#</th><th>Timestamp</th><th>IP Address</th><th>User Agent</th>"
            f"</tr>"
            f"</thead>"
            f"<tbody>{rows}</tbody>"
            f"</table>"
            f"</div>"
            f"</div>"
        )

        return format_html(visit_table)

    visit_details.short_description = 'Visit Details'

    def approval_history_table(self, obj):
        history = obj.approval_history.all().order_by('-action_date')
        if not history.exists():
            return "No approval history available."

        rows = ""
        for i, record in enumerate(history, start=1):
            rows += (
                f"<tr>"
                f"<td>{i}</td>"
                f"<td>{record.status}</td>"
                f"<td>{record.founder_member_amount}</td>"
                f"<td>{record.valid_date.strftime('%d-%m-%Y')}</td>"
                f"<td>{record.tsr_name}</td>"
                f"<td>{record.action_date.strftime('%d-%m-%Y')}</td>"
                f"</tr>"
            )

        approval_history_table = (
            f"<div style='flex: 1;'>"
            f"<h3>Approval History</h3>"
            f"<div style='max-height: 200px; overflow-y: auto;'>"
            f"<table style='border: 1px solid #ddd; border-collapse: collapse; width: 100%;'>"
            f"<thead>"
            f"<tr style='background-color: #f9f9f9;'>"
            f"<th>#</th><th>Status</th><th>Amount</th><th>Valid Till</th><th>TSR Name</th><th>Action Date</th>"
            f"</tr>"
            f"</thead>"
            f"<tbody>{rows}</tbody>"
            f"</table>"
            f"</div>"
            f"</div>"
        )

        return format_html(approval_history_table)

    approval_history_table.short_description = 'Approval History'

    def visit_and_approval_details(self, obj):
        visit_table = self.visit_details(obj)
        approval_table = self.approval_history_table(obj)

        combined_tables = (
            f"<div style='display: flex; flex-direction: row; justify-content: space-between;'>"
            f"{approval_table}"
            f"{visit_table}"
            f"</div>"
        )

        return format_html(combined_tables)

    visit_and_approval_details.short_description = 'Logs Details'

    approval_history_table.short_description = "Approval History"

    def valid_and_regdate(self, obj):
        return format_html(
            '<span style="">Reg:{}</span><br>'
            '<span style="color: red;">Valid:{}</span>',
            obj.regdate.strftime('%d-%m-%Y') if obj.regdate else ' - ',
            obj.valid_date.strftime('%d-%m-%Y') if obj.valid_date else ' - '
        )
    
    valid_and_regdate.short_description = 'Dates'

    @admin.display(description='Status')
    def status_display(self, obj):
        return "Active" if obj.is_active else "Inactive"

    def email_and_password(self, obj):
        return format_html(
            '<span style="">{}</span><br>'
            '<span style="color: red;">{}</span>',
            obj.email,
            obj.passwordd
        )
    
    email_and_password.short_description = 'Email & Password'

    def name_and_company(self, obj):
        return format_html(
            '<span style="">{}</span><br>'
            '<span style="color: red;">{}</span>',
            obj.name,
            obj.company
        )
    
    name_and_company.short_description = 'Name & Company'

    def address_detailed(self, obj):
        return format_html(
            '<span style="">{}</span><br>'
            '<span style="color: orange;">{}</span> '
            '<span style="color: red;">{}</span> ('
            '<span style="color: red;">{}</span>)',
            obj.address,
            obj.city,
            obj.state,
            obj.pin,
        )
    
    address_detailed.short_description = 'Address/City/State/Pin'

    def mobile_detailed(self, obj):
        return format_html(
            '<span style="">{}</span><br>'
            '<span style="color: red;">{}</span>',
            obj.mobileno,
            obj.telephone,
        )
    
    mobile_detailed.short_description = 'Mobile & Telephone'

    def founder_detail(self, obj):
        return format_html(
            '<span style="color:green;">{}</span><br>'
            '<span style="">Amount:{}</span><br>'
            '<span style="color: red;">TSR:{}</span>',
            obj.founder_member if obj.founder_member == 'Founder Member' else '',
            obj.founder_member_amount if obj.founder_member_amount else '-',
            obj.salesman.name if obj.salesman else '-'
        )

    
    founder_detail.short_description = 'Founder'

    def view_detail(self, obj):
        return format_html(
            '<span style="">:{}</span><br>'
            '<span style="color: red;">:{}</span>',
            obj.user_monthly_page_view if obj.user_monthly_page_view else '',
            obj.user_page_view_month if obj.user_page_view_month else '',
        )

    
    view_detail.short_description = 'Page'

    def formfield_for_foreignkey(self, db_field, request, **kwargs):
        if db_field.name == 'salesman':
            kwargs['queryset'] = User.objects.filter(is_salesman=True, is_active=True)
        return super().formfield_for_foreignkey(db_field, request, **kwargs)

    def save_model(self, request, obj, form, change):
        if not obj.username:
            obj.username = obj.email
        if not obj.email:
            obj.email = obj.username

        if change:
            item = ApprovalHistory.objects.create(
                user_id=obj.pk,
                status=obj.status,
                founder_member_amount=obj.founder_member_amount,
                valid_date=obj.valid_date,
                tsr_name=obj.salesman.name if obj.salesman else '',
                action_date=datetime.now(),
            )
            item.save()

        super().save_model(request, obj, form, change)


@admin.register(SalesmanUser)
class SalesmanAdmin(BaseUserAdmin):
    form = CustomUserChangeForm
    resource_class = UserResource
    add_form = CustomUserCreationForm
    change_password_form = AdminPasswordChangeForm
    filter_horizontal = ('user_permissions',)

    list_display = ('email_and_password', 'name_and_company', 'mobile_detailed', 'address_detailed', 'status_display')
    list_display_links = ('email_and_password', 'name_and_company',)
    search_fields = ('email', 'name', 'city', 'company', 'mobileno')
    ordering = ('-pk',)
    readonly_fields = ('passwordd', 'regdate', 'last_login', 'is_login',)
    sortable_by = ('-pk', 'state', 'city', 'email')

    fieldsets = (
        (_('Personal Info'), {
            'fields': (
                ('email', 'username', 'name', 'company', 'address', 'city', 'state', 'pin', 
                 'telephone', 'mobileno', 'fax', 'status', 'is_active', 'is_staff', 'is_salesman', 'is_superuser'),
                ('valid_date',),
                ('passwordd', 'regdate', 'last_login', 'is_login',),
                ('user_permissions'),
            ),
        }),
    )

    @admin.display(description='Status')
    def status_display(self, obj):
        return "Active" if obj.is_active else "Inactive"

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        return qs.filter(Q(is_salesman=True) | Q(is_staff=True) | Q(is_superuser=True))

    def get_actions(self, request):
        actions = super().get_actions(request)
        if "delete_selected" in actions:
            del actions["delete_selected"]
        return actions

    def email_and_password(self, obj):
        return format_html(
            '<span style="">{}</span><br>'
            '<span style="color: red;">{}</span>',
            obj.email,
            obj.passwordd
        )
    
    email_and_password.short_description = 'Email & Password'

    def name_and_company(self, obj):
        return format_html(
            '<span style="">{}</span><br>'
            '<span style="color: red;">{}</span>',
            obj.name,
            obj.company
        )
    
    name_and_company.short_description = 'Name & Company'

    def address_detailed(self, obj):
        return format_html(
            '<span style="">{}</span><br>'
            '<span style="color: red;">{}</span> '
            '<span style="color: red;">{}</span> ('
            '<span style="color: red;">{}</span>)',
            obj.address,
            obj.city,
            obj.state,
            obj.pin,
        )
    
    address_detailed.short_description = 'Address/City/State/Pin'

    def mobile_detailed(self, obj):
        return format_html(
            '<span style="">{}</span><br>'
            '<span style="color: red;">{}</span>',
            obj.mobileno,
            obj.telephone,
        )
    
    mobile_detailed.short_description = 'Mobile & Telephone'

    def save_model(self, request, obj, form, change):
        if not obj.username: obj.username = obj.email
        if not obj.email: obj.email = obj.username
        super().save_model(request, obj, form, change)

admin.site.unregister(UserVisit)

@admin.register(UserVisit)
class UserVisitAdmin(admin.ModelAdmin):
    list_display = ('timestamp', 'user', 'session_key', 'remote_addr', 'ua_string')
    list_display_links = list_display
    list_filter = (('timestamp', admin.DateFieldListFilter), 'user',)
    search_fields = ('user__name', 'remote_addr', )

@admin.register(AuthSession)
class AuthSessionAdmin(admin.ModelAdmin):
    list_display = ('user', 'ip_address', 'user_agent', 'created_at', 'invalidated_at', 'status_display',)
    list_filter = ('user', 'is_active', 'created_at', 'invalidated_at')
    search_fields = ('user__username', 'ip_address', 'token_jti')
    readonly_fields = ('created_at', 'invalidated_at')
    fieldsets = (
        (None, {'fields': (
            ('user',),
            ('token_jti', 'user_agent'),
            ('ip_address', 'is_active'),
            ('created_at', 'invalidated_at', 'invalidation_reason')
        )}),
    )

    @admin.display(description='Status')
    def status_display(self, obj):
        return "Active" if obj.is_active else "Inactive"